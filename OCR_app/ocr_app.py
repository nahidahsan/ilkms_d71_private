import easyocr as ocr  #OCR
import streamlit as st  #Web App
from PIL import Image #Image Processing
import numpy as np #Image Processing 
import os, io
from google.cloud import vision
from google.cloud.vision_v1 import types
import cv2
import streamlit.components.v1 as components
from streamlit_ace import st_ace



header="""
<head>
<style>

.header{
left: 0;
width: 100%;
background-color: white;
color: black;
text-align: center;
margin:0;
}


</style>
</head>
<body>
<div class="header">
<p style='text-align: center'><img src="https://cdn.worldvectorlogo.com/logos/bangladesh-govt-logo.svg" width = 16% height = 16%></p>
</div>
</body>
"""

st.markdown(header, unsafe_allow_html=True)

st.markdown("<h5 style='text-align: center; color: black; font-weight: bold;'>Intelligent Land Knowledge Management System <br>for Ministry of Land</h5>", unsafe_allow_html=True)


#image uploader

col1, col2 = st.columns(2)
uploaded_file = st.file_uploader("Choose a image file", type="jpg", accept_multiple_files=True)
images = []
if uploaded_file is not None:
    for file in uploaded_file:
        # Convert the file to an opencv image.
        file_bytes = np.asarray(bytearray(file.read()), dtype=np.uint8)
        opencv_image = cv2.imdecode(file_bytes, 1)
        images.append(opencv_image)
    # Now do something with the image! For example, let's display it:
col1.image(images,use_column_width=True)


@st.cache
def image_process(image, delta = 1, limit = 5):

#         gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        inverted_image = cv2.bitwise_not(image)
        binary_image = cv2.cvtColor(inverted_image, cv2.COLOR_BGR2GRAY)
        thresh, im_bw = cv2.threshold(binary_image,90,255, cv2.THRESH_BINARY)
        final_image = cv2.bitwise_not(im_bw)

        # correct skewness of image
        image = final_image
        def determine_score(arr, angle):
            data = inter.rotate(arr, angle, reshape=False, order=0)
            histogram = np.sum(data, axis=1, dtype=float)
            score = np.sum((histogram[1:] - histogram[:-1]) ** 2, dtype=float)
            return histogram, score

    #     gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        thresh = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1] 

        scores = []
        angles = np.arange(-limit, limit + delta, delta)
        for angle in angles:
            histogram, score = determine_score(thresh, angle)
            scores.append(score)

        best_angle = angles[scores.index(max(scores))]

        (h, w) = image.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, best_angle, 1.0)
        rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

        return rotated
#     processed_image = [image_process(image) for image in opencv_image]
#     col2.image(opencv_image)

def detect_text(processed_image):
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r'VisionApi_token.json'
    client = vision.ImageAnnotatorClient()
    def decode_png(image_data):
        success, encoded_image = cv2.imencode('.jpg', image_data)
        content2 = encoded_image.tobytes()
        return content2 

    decode_image = [decode_png(image) for image in processed_image]
    content_image = [types.Image(content = image) for image in decode_image]

    bangla_text = dict()
    for idx, image in enumerate(content_image):
        response = client.text_detection(image = image)
        texts = response.text_annotations

        for text in texts:
            bangla_text[idx] = ('\n"{}"'.format(text.description))
            break

    for key, value in bangla_text.items():
        print(value)

    return bangla_text


    #editing panel & save document
from docx import Document
from docx.shared import Inches

document = Document()

# document.add_heading('জমিজমার আইন', 0)


count = 0
bangla_text = detect_text(images)

for key, value in bangla_text.items():
    txt = col2.text_area(label ="",value=value, height =500)  
    document.add_paragraph(txt)
    document.add_page_break()

    

# for key, value in bangla_text.items():
#     document.add_paragraph(value)
#     document.add_page_break()
    
# p = document.add_paragraph()
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

document.save('demo.docx')

from docx2pdf import convert
import pythoncom
import win32com.client
xl=win32com.client.Dispatch("Excel.Application",pythoncom.CoInitialize())
c =  st.container()
c1, c2, c3 = c.columns(3)




with open("demo.docx", "rb") as docx_file:
    docbyte = docx_file.read()
    convert("demo.docx", "output.pdf") 


c2.download_button(label="Download DocX",
                    data=docbyte,
                    file_name="test.docx",
                    mime='application/octet-stream')

with open("output.pdf", "rb") as pdf_file:
    PDFbyte = pdf_file.read()

c2.download_button(label="Download PDF",
                    data=PDFbyte,
                    file_name="output.pdf",
                    mime='application/octet-stream')
    


# hide header and footer
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)



footer="""<style>
a:link , a:visited{
color: blue;
background-color: transparent;
text-decoration: underline;
}

a:hover,  a:active {
color: red;
background-color: transparent;
text-decoration: underline;
}

.footer {
position: fixed;
left: 1%;
bottom: 0;
width: 98%;
background-color: white;
color: black;
text-align: center;
overflow:auto;

margin:0;
}
.footerP{
color:black;
font-family:FiraSans-Regular, sans-serif;
font-size:12px;
margin:0;
line-height:22.5vh;
}

.socialMedias{
float:right;
font-family:FiraSans-Regular, sans-serif;
font-size:12px;
}
.socialMedias2{
float:left;
font-family:FiraSans-Regular, sans-serif;
font-size:12px
}
.facebook{
width:5vw;
height:10vh;
}
</style>
<div class="footer">
<p style='text-align: right' class = "socialMedias">Design and Developed by <img src="https://imgs.search.brave.com/I5o3znS3QqZZYU2R7W7aP15HHbIIyPCqfGFmtp_TvZ4/rs:fit:706:225:1/g:ce/aHR0cHM6Ly90c2Ux/LmV4cGxpY2l0LmJp/bmcubmV0L3RoP2lk/PU9JUC5IWUtRXzRI/ME8xZ3k4Sko1ME85/aDRBQUFBQSZwaWQ9/QXBp" width = 12% height = 12%></p>
<p style='text-align: left' class = "socialMedias2">Copyrights © 2022, Ministry of Land, All Rights Reserved</p>
</div>
"""

st.markdown(footer, unsafe_allow_html=True)